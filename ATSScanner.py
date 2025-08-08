'''import os
import streamlit as st
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains import LLMChain
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings  
from langchain_core.runnables import RunnableLambda, RunnableMap
import google.generativeai as genai
from dotenv import load_dotenv
import shutil
import re
from docx import Document
from io import BytesIO

# Load environment variables
load_dotenv()

# Configure Google AI API
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    st.error("Please set your GOOGLE_API_KEY in a .env file")
    st.stop()

genai.configure(api_key=GOOGLE_API_KEY)

# Setup embedding model
embedding_model = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Create or load Chroma vector store
VECTOR_STORE_DIR = "chroma_store"
if os.path.exists(VECTOR_STORE_DIR):
    vectorstore = Chroma(persist_directory=VECTOR_STORE_DIR, embedding_function=embedding_model)
else:
    os.makedirs(VECTOR_STORE_DIR, exist_ok=True)
    vectorstore = Chroma(persist_directory=VECTOR_STORE_DIR, embedding_function=embedding_model)

# Extract text from uploaded files
def extract_text_from_resume(file):
    temp_file_path = f"temp_{file.name}"
    with open(temp_file_path, "wb") as f:
        f.write(file.getbuffer())

    file_extension = os.path.splitext(file.name)[1].lower()
    try:
        if file_extension == '.pdf':
            loader = PyPDFLoader(temp_file_path)
        elif file_extension == '.docx':
            loader = Docx2txtLoader(temp_file_path)
        elif file_extension == '.txt':
            loader = TextLoader(temp_file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        documents = loader.load()
        text = " ".join([doc.page_content for doc in documents])
        return text
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

# Text splitting
def split_text(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    return splitter.create_documents([text])

# Store resume analysis in vector store
def store_resume_analysis(resume_text, analysis, doc_id):
    documents = split_text(analysis)
    vectorstore.add_documents(documents, ids=[f"{doc_id}_chunk_{i}" for i in range(len(documents))])
    vectorstore.persist()

# Extract percentage score from analysis text
def extract_suitability_score(text):
    match = re.search(r"Suitability Score: (\d{1,3})%", text)
    if match:
        return int(match.group(1))
    return None

# Create updated resume as Word document
def create_updated_resume_doc(text):
    doc = Document()
    doc.add_heading('Updated Resume (ATS Optimized)', 0)
    for line in text.split("\n"):
        doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Main App
def main():
    st.set_page_config(page_title="Resume Screener", layout="wide")
    st.title("🎯 Resume Screening + ATS Optimizer")

    col1, col2 = st.columns(2)
    with col1:
        st.header("📝 Job Description")
        job_requirements = st.text_area("Paste the job description here", height=300)
    with col2:
        st.header("📄 Upload Resume")
        uploaded_file = st.file_uploader("Upload your resume", type=["pdf", "docx", "txt"])

    if st.button("🚀 Analyze Resume") and uploaded_file and job_requirements:
        with st.spinner("Processing..."):
            resume_text = extract_text_from_resume(uploaded_file)
            st.subheader("📃 Extracted Resume Text")
            st.text_area("View Resume Text", resume_text, height=300)

            llm = ChatGoogleGenerativeAI(
                model="gemini-2.0-flash",
                google_api_key=GOOGLE_API_KEY,
                temperature=0.2
            )

            # Prompt for analysis
            prompt_analysis = PromptTemplate(
                input_variables=["job_requirements", "resume_text"],
                template="""
You are an expert ATS resume screener. Analyze the resume against the job requirements.

1. Extract all job titles with company names.
2. Identify important job requirements missing in the resume.
3. Suggest improvements.
4. Provide ATS Suitability Score (0–100%).

Resume:
{resume_text}

Job Requirements:
{job_requirements}

Format strictly:
### Job Titles and Companies:
<list>

### Missing Points in Resume:
<list>

### ATS Suitability Score:
XX%

### Suggestions:
<text>
"""
            )

            analysis_chain = (
                RunnableMap({
                    "job_requirements": lambda x: x["job_requirements"],
                    "resume_text": lambda x: x["resume_text"]
                }) | prompt_analysis | llm | StrOutputParser()
            )

            analysis = analysis_chain.invoke({
                "job_requirements": job_requirements,
                "resume_text": resume_text
            })

            st.subheader("📊 AI Screening Result")
            st.markdown(analysis)

            score = extract_suitability_score(analysis)
            if score is not None:
                st.metric("ATS Suitability Score", f"{score}%")

            store_resume_analysis(resume_text, analysis, os.path.splitext(uploaded_file.name)[0])
            st.success("✅ Analysis stored in vector database.")

            st.download_button("⬇️ Download Analysis", analysis, file_name="resume_analysis.txt")

            # Generate updated resume
            st.subheader("📄 Generate Updated ATS Resume")

            prompt_updated_resume = PromptTemplate(
                input_variables=["job_requirements", "resume_text"],
                template="""
You are an expert ATS resume writer. Rewrite the candidate’s resume to meet 96%+ ATS match to the job description below.

Resume:
{resume_text}

Job Description:
{job_requirements}

Your output should be a professional, detailed, ATS-optimized resume that includes missing keywords, relevant achievements, and matches job title and responsibilities. DO NOT include personal details.
"""
            )

            resume_update_chain = (
                RunnableMap({
                    "job_requirements": lambda x: x["job_requirements"],
                    "resume_text": lambda x: x["resume_text"]
                }) | prompt_updated_resume | llm | StrOutputParser()
            )

            updated_resume = resume_update_chain.invoke({
                "job_requirements": job_requirements,
                "resume_text": resume_text
            })

            docx_data = create_updated_resume_doc(updated_resume)
            st.download_button(
                label="📥 Download Updated Resume (Word)",
                data=docx_data,
                file_name="Updated_ATS_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
'''
import os
import streamlit as st
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains import LLMChain
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.runnables import RunnableLambda, RunnableMap
import google.generativeai as genai
from dotenv import load_dotenv
import re
from docx import Document
from io import BytesIO

# Load environment variables
load_dotenv()

# Configure Google AI API
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    st.error("Please set your GOOGLE_API_KEY in a .env file")
    st.stop()

genai.configure(api_key=GOOGLE_API_KEY)

# Extract text from uploaded files
def extract_text_from_resume(file):
    file_extension = os.path.splitext(file.name)[1].lower()
    
    try:
        if file_extension == '.pdf':
            with open("temp.pdf", "wb") as f:
                f.write(file.getbuffer())
            loader = PyPDFLoader("temp.pdf")
        elif file_extension == '.docx':
            with open("temp.docx", "wb") as f:
                f.write(file.getbuffer())
            loader = Docx2txtLoader("temp.docx")
        elif file_extension == '.txt':
            with open("temp.txt", "wb") as f:
                f.write(file.getbuffer())
            loader = TextLoader("temp.txt")
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

        documents = loader.load()
        text = " ".join([doc.page_content for doc in documents])
        return text
    finally:
        if os.path.exists("temp.pdf"): os.remove("temp.pdf")
        if os.path.exists("temp.docx"): os.remove("temp.docx")
        if os.path.exists("temp.txt"): os.remove("temp.txt")

# Extract percentage score from analysis text
def extract_suitability_score(text):
    match = re.search(r"ATS Suitability Score: (\d{1,3})%", text)
    if match:
        return int(match.group(1))
    return None

# Create updated resume as a Word document
def create_updated_resume_doc(text):
    doc = Document()
    doc.add_heading('Updated Resume (ATS Optimized)', 0)
    
    lines = text.split('\n')
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Main App
def main():
    st.set_page_config(page_title="Resume Screener", layout="wide")
    st.title("🎯 Resume Screening + ATS Optimizer")

    col1, col2 = st.columns(2)
    with col1:
        st.header("📝 Job Description")
        job_requirements = st.text_area("Paste the job description here", height=300)
    with col2:
        st.header("📄 Upload Resume")
        uploaded_file = st.file_uploader("Upload your resume", type=["pdf", "docx", "txt"])

    if st.button("🚀 Analyze Resume") and uploaded_file and job_requirements:
        with st.spinner("Processing..."):
            resume_text = extract_text_from_resume(uploaded_file)
            st.subheader("📃 Extracted Resume Text")
            st.text_area("View Resume Text", resume_text, height=300)

            llm = ChatGoogleGenerativeAI(
                model="gemini-2.0-flash",
                google_api_key=GOOGLE_API_KEY,
                temperature=0.2
            )

            # Prompt for analysis
            prompt_analysis = PromptTemplate(
                input_variables=["job_requirements", "resume_text"],
                template="""
You are an expert ATS resume screener. Analyze the resume against the job requirements.

1. Extract all job titles with company names.
2. Identify important job requirements missing in the resume.
3. Suggest improvements.
4. Provide ATS Suitability Score (0–100%).

Resume:
{resume_text}

Job Requirements:
{job_requirements}

Format strictly:
### Job Titles and Companies:
<list>

### Missing Points in Resume:
<list>

### ATS Suitability Score:
XX%

### Suggestions:
<text>
"""
            )

            analysis_chain = (
                RunnableMap({
                    "job_requirements": lambda x: x["job_requirements"],
                    "resume_text": lambda x: x["resume_text"]
                }) | prompt_analysis | llm | StrOutputParser()
            )

            analysis = analysis_chain.invoke({
                "job_requirements": job_requirements,
                "resume_text": resume_text
            })

            st.subheader("📊 AI Screening Result")
            st.markdown(analysis)

            score = extract_suitability_score(analysis)
            if score is not None:
                st.metric("ATS Suitability Score", f"{score}%")

            st.download_button("⬇️ Download Analysis", analysis, file_name="resume_analysis.txt")

            # Generate updated resume
            st.subheader("📄 Generate Updated ATS Resume")

            with st.spinner("Generating updated resume..."):
                prompt_updated_resume = PromptTemplate(
                    input_variables=["job_requirements", "resume_text"],
                    template="""
You are an expert ATS resume writer. Rewrite the candidate’s resume to meet 96%+ ATS match to the job description below.

Resume:
{resume_text}

Job Description:
{job_requirements}

Your output should be a professional, detailed, ATS-optimized resume that includes missing keywords, relevant achievements, and matches job title and responsibilities. DO NOT include personal details.
"""
                )

                resume_update_chain = (
                    RunnableMap({
                        "job_requirements": lambda x: x["job_requirements"],
                        "resume_text": lambda x: x["resume_text"]
                    }) | prompt_updated_resume | llm | StrOutputParser()
                )

                updated_resume = resume_update_chain.invoke({
                    "job_requirements": job_requirements,
                    "resume_text": resume_text
                })

            docx_data = create_updated_resume_doc(updated_resume)
            st.download_button(
                label="📥 Download Updated Resume (Word)",
                data=docx_data,
                file_name="Updated_ATS_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()